from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUT_FILE = BASE_DIR / "Nutrition_Grade_Prediction_Detailed_Report.docx"
FIG_DIR = BASE_DIR / "ieee_report" / "figures"
WORD_FIG_DIR = BASE_DIR / "word_report_figures"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], LIGHT_BLUE)
        if widths:
            hdr[i].width = Inches(widths[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 25 else WD_ALIGN_PARAGRAPH.CENTER
            size = 8 if len(str(value)) > 80 else 9
            set_cell_text(cells[i], value, align=align, size=size)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)


def add_figure(doc, filename, caption, width=5.9):
    path = FIG_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_word_figure(doc, filename, caption, width=5.9):
    path = WORD_FIG_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.2

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    return doc


def build_report():
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Machine Learning and Deep Learning-Based Nutrition Grade Prediction")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Detailed Final Project Report")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(80, 80, 80)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Dataset: Open Food Facts | Target: nutrition_grade_fr | Task: Multi-Class Classification")
    doc.add_paragraph()

    add_heading(doc, "Abstract", 1)
    add_body(
        doc,
        "This project focuses on predicting food nutrition grades using the Open Food Facts dataset. "
        "The original dataset was cleaned, reduced to 20 meaningful nutritional and ingredient-based "
        "features, and prepared for multi-class classification. The final dataset contains 14,070 "
        "records, no missing values, no duplicate records, and five target classes. Five machine "
        "learning models and three deep learning models were implemented and compared using accuracy, "
        "precision, recall, macro F1-score, and confusion matrix analysis. XGBoost achieved the best "
        "full-feature performance, while the top-10 feature XGBoost model achieved the highest overall "
        "accuracy. The project also included data standardization, class balancing with SMOTE, feature "
        "selection, training-time analysis, and model serialization for future deployment."
    )
    add_body(
        doc,
        "Keywords: Open Food Facts, nutrition grade prediction, machine learning, deep learning, "
        "XGBoost, SMOTE, feature selection, food classification."
    )

    add_heading(doc, "1. Introduction", 1)
    add_body(
        doc,
        "Food products contain several nutritional indicators such as energy, fat, sugar, salt, fiber, "
        "protein, vitamins, minerals, additives, and ingredient information. Although these values are "
        "available on nutrition labels, manually interpreting them can be difficult for consumers. "
        "Automated nutrition grade prediction can support healthier food decisions and provide a data-driven "
        "way to evaluate product quality."
    )
    add_body(
        doc,
        "The purpose of this project is to predict the nutrition grade of food products from selected "
        "Open Food Facts attributes. The target variable is nutrition_grade_fr, which contains five classes: "
        "a, b, c, d, and e. In the implemented notebook, the classes were ordinally encoded as a=4, b=3, "
        "c=2, d=1, and e=0. Since the model predicts one of five possible classes, the task is a multi-class "
        "classification problem."
    )

    add_heading(doc, "2. Related Work", 1)
    add_body(
        doc,
        "Machine learning has been widely used for structured health and nutrition-related classification "
        "tasks. Linear models such as Logistic Regression are useful baselines, but they may fail to capture "
        "nonlinear relationships among nutritional features. Tree-based models and ensemble methods often "
        "perform better on tabular data because they can model interactions between features without requiring "
        "feature scaling."
    )
    add_body(
        doc,
        "Gradient boosting methods such as XGBoost are especially strong for structured tabular datasets. "
        "Deep learning models such as MLP, DNN, and 1D CNN can also learn nonlinear patterns, but for relatively "
        "small tabular datasets, tree-based ensemble methods frequently outperform neural networks. This project "
        "therefore compares both traditional machine learning and deep learning approaches."
    )

    add_heading(doc, "3. Methodology", 1)
    add_heading(doc, "3.1 Data Collection", 2)
    add_body(
        doc,
        "The dataset was derived from Open Food Facts, a large public food product database. The raw dataset "
        "contained many columns and many incomplete records. A subset of 20 nutritionally relevant features "
        "was selected to satisfy the project requirement while keeping the dataset meaningful for nutrition "
        "grade prediction."
    )

    add_heading(doc, "3.2 Dataset Overview", 2)
    add_table(
        doc,
        ["Property", "Value"],
        [
            ["Number of records", "14,070"],
            ["Number of features", "20"],
            ["Target variable", "nutrition_grade_fr"],
            ["Target classes", "5"],
            ["Missing values", "0"],
            ["Duplicate records", "0"],
        ],
        widths=[3.0, 3.0],
    )
    add_table(
        doc,
        ["Class", "Count", "Percentage"],
        [
            ["a", "3,188", "22.66%"],
            ["b", "3,337", "23.72%"],
            ["c", "3,142", "22.33%"],
            ["d", "3,297", "23.43%"],
            ["e", "1,106", "7.86%"],
        ],
        widths=[1.5, 2.0, 2.0],
    )
    add_figure(doc, "target_distribution.png", "Figure 1. Target class distribution.", width=5.5)

    add_heading(doc, "3.3 Data Preprocessing", 2)
    add_body(
        doc,
        "The preprocessing stage included missing value handling, duplicate removal, zero-value inspection, "
        "target encoding, and standardization for scale-sensitive models."
    )
    add_heading(doc, "3.3.1 Handling Missing Values", 3)
    add_body(
        doc,
        "Records with missing values in the selected 20 features or target variable were removed. The final "
        "dataset contains zero missing values, so no imputation method was required."
    )
    add_heading(doc, "3.3.2 Duplicate Removal", 3)
    add_body(
        doc,
        "Duplicate records were removed to avoid inflated model performance and reduce the risk of data leakage. "
        "The final dataset contains 14,070 unique records."
    )
    add_heading(doc, "3.3.3 Zero Value Analysis", 3)
    add_body(
        doc,
        "Some features naturally contain many zero values. For example, most products do not contain trans fat "
        "or palm oil, so zero values in these columns can be meaningful rather than missing. However, records "
        "with effectively empty nutrition tables were removed."
    )
    add_figure(doc, "zero_value_distribution.png", "Figure 2. Zero value percentage by feature.", width=5.8)
    add_heading(doc, "3.3.4 Standardization", 3)
    add_body(
        doc,
        "StandardScaler was applied to Logistic Regression, KNN, and deep learning models because these models "
        "are sensitive to feature scale. The scaler was fitted only on the training data and then applied to the "
        "test data to prevent data leakage."
    )

    add_heading(doc, "3.4 Feature Analysis and Correlation", 2)
    add_body(
        doc,
        "The selected features cover energy, macronutrients, salt-related features, micronutrients, additives, "
        "palm oil indicators, and ingredient count. Correlation analysis showed that salt_100g and sodium_100g "
        "are perfectly correlated because salt and sodium are mathematically related. Both were kept because the "
        "project required exactly 20 features and both are nutritionally meaningful."
    )
    add_table(
        doc,
        ["Feature 1", "Feature 2", "Correlation"],
        [
            ["salt_100g", "sodium_100g", "1.0000"],
            ["energy_100g", "fat_100g", "0.7888"],
            ["additives_n", "ingredients_count", "0.7467"],
            ["energy_100g", "carbohydrates_100g", "0.7103"],
            ["fat_100g", "saturated-fat_100g", "0.6692"],
            ["carbohydrates_100g", "sugars_100g", "0.5739"],
        ],
        widths=[2.3, 2.3, 1.2],
    )
    add_figure(doc, "feature_correlation_heatmap.png", "Figure 3. Feature correlation heatmap.", width=5.9)

    add_heading(doc, "4. Model Development", 1)
    add_body(
        doc,
        "The project trained five machine learning models and three deep learning models. The machine learning "
        "models were Logistic Regression, KNN, Decision Tree, Random Forest, and XGBoost. The deep learning "
        "models were MLP, DNN, and 1D CNN."
    )
    add_table(
        doc,
        ["Sr No", "Model", "Tuning Parameters"],
        [
            [
                "1",
                "Logistic Regression",
                "max_iter=5000, class_weight='balanced', random_state=42; StandardScaler applied",
            ],
            [
                "2",
                "K-Nearest Neighbors",
                "GridSearchCV(cv=5, scoring='f1_macro'); best: n_neighbors=5, weights='distance', metric='manhattan'",
            ],
            [
                "3",
                "Decision Tree",
                "criterion='entropy', max_depth=12, min_samples_leaf=5, min_samples_split=2, class_weight='balanced', random_state=42",
            ],
            [
                "4",
                "Random Forest",
                "n_estimators=200, criterion='entropy', max_depth=12, min_samples_split=5, min_samples_leaf=10, class_weight='balanced', random_state=42, n_jobs=-1",
            ],
            [
                "5",
                "XGBoost",
                "n_estimators=300, max_depth=5, learning_rate=0.05, subsample=0.8, colsample_bytree=0.8, objective='multi:softprob', eval_metric='mlogloss', random_state=42, n_jobs=-1",
            ],
            [
                "6",
                "MLP",
                "Dense(128)-Dropout(0.30)-Dense(64)-Dropout(0.20)-Dense(32)-Softmax(5), optimizer='adam', loss='categorical_crossentropy', batch_size=32, epochs=100, early_stopping patience=10",
            ],
            [
                "7",
                "DNN",
                "Dense(256)-BatchNorm-Dropout(0.35)-Dense(128)-BatchNorm-Dropout(0.30)-Dense(64)-BatchNorm-Dropout(0.20)-Dense(32)-Softmax(5), optimizer='adam', epochs=120",
            ],
            [
                "8",
                "1D CNN",
                "Input shape=(20,1), Conv1D(64,kernel=3)-BatchNorm-MaxPool-Dropout(0.25), Conv1D(128,kernel=3)-BatchNorm-MaxPool-Dropout(0.30), Dense(64), Softmax(5), epochs=120",
            ],
        ],
        widths=[0.65, 1.65, 4.2],
    )
    add_heading(doc, "4.1 Machine Learning Models", 2)
    add_bullets(
        doc,
        [
            "Logistic Regression: used as a linear baseline with StandardScaler, max_iter=5000, and class_weight=balanced.",
            "KNN: optimized using 5-fold GridSearchCV with f1_macro scoring; best parameters were metric=manhattan, n_neighbors=5, weights=distance.",
            "Decision Tree: trained with entropy, max_depth=12, min_samples_leaf=5, min_samples_split=2, and class_weight=balanced.",
            "Random Forest: trained with 200 estimators, entropy criterion, max_depth=12, min_samples_split=5, min_samples_leaf=10, and class_weight=balanced.",
            "XGBoost: trained with 300 estimators, max_depth=5, learning_rate=0.05, subsample=0.8, colsample_bytree=0.8, and balanced sample weights.",
        ],
    )

    add_heading(doc, "4.2 Deep Learning Models", 2)
    add_bullets(
        doc,
        [
            "MLP: dense neural network with ReLU activations, dropout, softmax output, early stopping, and class weights.",
            "DNN: deeper dense network with Batch Normalization and Dropout layers.",
            "1D CNN: convolutional model using reshaped feature vectors with input shape (20, 1).",
        ],
    )

    add_heading(doc, "5. Model Evaluation", 1)
    add_body(
        doc,
        "All models were evaluated using accuracy, macro precision, macro recall, and macro F1-score. Macro F1-score "
        "was important because the target classes are not perfectly balanced, especially class e."
    )
    add_table(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "Macro F1"],
        [
            ["Logistic Regression", "0.7591", "0.7574", "0.7763", "0.7642"],
            ["KNN with CV", "0.8547", "0.8608", "0.8514", "0.8556"],
            ["Decision Tree", "0.8834", "0.8770", "0.8863", "0.8812"],
            ["Random Forest", "0.8842", "0.8843", "0.8860", "0.8849"],
            ["XGBoost", "0.9325", "0.9330", "0.9308", "0.9318"],
            ["MLP", "0.9112", "0.9083", "0.9166", "0.9115"],
            ["DNN", "0.8952", "0.8919", "0.9022", "0.8962"],
            ["1D CNN", "0.8923", "0.8928", "0.8914", "0.8919"],
        ],
        widths=[2.1, 1.1, 1.1, 1.1, 1.1],
    )
    add_body(
        doc,
        "XGBoost achieved the best full-feature model performance with 0.9325 accuracy and 0.9318 macro F1-score. "
        "Among the deep learning models, MLP achieved the best performance with 0.9112 accuracy and 0.9115 macro F1-score."
    )

    add_heading(doc, "5.1 Confusion Matrix Analysis", 2)
    add_body(
        doc,
        "Confusion matrices were used to inspect class-level model behavior. They show how many samples from each "
        "actual nutrition grade were assigned to each predicted grade. Since the target variable has five classes, "
        "the confusion matrices are useful for identifying whether models mostly confuse neighboring grades or make "
        "larger classification errors."
    )
    add_word_figure(doc, "logistic_regression_confusion_matrix.png", "Figure 5. Logistic Regression confusion matrix.", width=5.2)
    add_word_figure(doc, "knn_confusion_matrix.png", "Figure 6. KNN confusion matrix.", width=5.2)
    add_word_figure(doc, "decision_tree_confusion_matrix.png", "Figure 7. Decision Tree confusion matrix.", width=5.2)
    add_word_figure(doc, "random_forest_confusion_matrix.png", "Figure 8. Random Forest confusion matrix.", width=5.2)
    add_word_figure(doc, "xgboost_confusion_matrix.png", "Figure 9. XGBoost confusion matrix.", width=5.2)
    add_word_figure(doc, "mlp_confusion_matrix.png", "Figure 10. MLP confusion matrix.", width=5.2)
    add_word_figure(doc, "dnn_confusion_matrix.png", "Figure 11. DNN confusion matrix.", width=5.2)
    add_word_figure(doc, "cnn_1d_confusion_matrix.png", "Figure 12. 1D CNN confusion matrix.", width=5.2)

    add_heading(doc, "6. SMOTE-Based Data Balancing", 1)
    add_body(
        doc,
        "SMOTE was applied to the best full-feature model, XGBoost, to balance the minority class. SMOTE was applied "
        "only to the training data. The test data was not modified and remained real data."
    )
    add_table(
        doc,
        ["Setting", "Accuracy", "Precision", "Recall", "Macro F1"],
        [
            ["Before SMOTE", "0.9325", "0.9330", "0.9308", "0.9318"],
            ["After SMOTE", "0.9314", "0.9316", "0.9299", "0.9306"],
        ],
        widths=[2.0, 1.1, 1.1, 1.1, 1.1],
    )
    add_body(
        doc,
        "The SMOTE model produced a very similar result to the original XGBoost model. This indicates that XGBoost "
        "with balanced sample weights already handled the class imbalance effectively."
    )
    add_word_figure(
        doc,
        "xgboost_smote_confusion_comparison.png",
        "Figure 13. XGBoost confusion matrices before and after SMOTE.",
        width=6.2,
    )

    add_heading(doc, "7. Feature Selection and Reduction", 1)
    add_body(
        doc,
        "Feature selection was performed using XGBoost feature importance scores. The top 10 features were selected "
        "and the XGBoost model was retrained using only these features."
    )
    add_table(
        doc,
        ["Feature Set", "No. Features", "Accuracy", "Recall", "Macro F1"],
        [
            ["All features", "20", "0.9325", "0.9308", "0.9318"],
            ["Top 10 features", "10", "0.9350", "0.9340", "0.9340"],
        ],
        widths=[2.0, 1.3, 1.0, 1.0, 1.0],
    )
    add_body(
        doc,
        "The top-10 feature model slightly outperformed the full-feature XGBoost model. This suggests that some "
        "features may be redundant or low-signal for this task."
    )
    add_word_figure(
        doc,
        "xgboost_top10_confusion_matrix.png",
        "Figure 14. XGBoost Top-10 feature confusion matrix.",
        width=5.2,
    )
    add_figure(doc, "xgboost_top10_feature_importance.png", "Figure 4. Top 10 XGBoost feature importances.", width=5.6)

    add_heading(doc, "8. Training Time and Runtime Analysis", 1)
    add_body(
        doc,
        "Training time and one-record prediction runtime were measured using the same 80%-20% train-test split. "
        "The top-10 XGBoost model trained faster than the full-feature XGBoost model while also achieving higher accuracy."
    )
    add_table(
        doc,
        ["Model", "Training Time (s)", "One Record Runtime (s)"],
        [
            ["Logistic Regression", "0.7404", "0.000156"],
            ["KNN Optimized", "0.0027", "0.001129"],
            ["Decision Tree", "0.1654", "0.001031"],
            ["Random Forest", "4.9080", "0.055375"],
            ["XGBoost Full", "4.8410", "0.004218"],
            ["XGBoost Top 10", "2.5713", "0.001816"],
            ["MLP", "60.7906", "0.296751"],
            ["DNN", "40.8359", "0.392009"],
            ["1D CNN", "43.9195", "0.442575"],
        ],
        widths=[2.3, 2.0, 2.0],
    )

    add_heading(doc, "9. Discussion", 1)
    add_body(
        doc,
        "The results show that XGBoost is the most successful full-feature model for this dataset. This is expected "
        "because Open Food Facts data is structured tabular data, and gradient boosting algorithms are highly effective "
        "for this type of problem. The top-10 feature XGBoost model achieved the highest accuracy, showing that feature "
        "selection improved the model slightly."
    )
    add_body(
        doc,
        "Deep learning models performed well but did not outperform XGBoost. The MLP model was the best deep learning "
        "model, while DNN and 1D CNN produced slightly lower results. This is reasonable because deep learning models "
        "often need larger and more complex datasets to outperform tree-based methods on tabular data."
    )
    add_body(
        doc,
        "The SMOTE experiment showed that class balancing did not improve XGBoost performance. This may be because "
        "sample weighting already reduced the effect of class imbalance, and the minority class was not extremely small "
        "after preprocessing."
    )

    add_heading(doc, "10. Conclusion", 1)
    add_body(
        doc,
        "This project built a complete machine learning and deep learning pipeline for predicting nutrition grades from "
        "Open Food Facts data. The final dataset contained 14,070 records, 20 features, no missing values, and no duplicates. "
        "Eight models were trained and evaluated. XGBoost achieved the best full-feature performance, while the top-10 "
        "feature XGBoost model achieved the best overall accuracy of 0.9350."
    )
    add_body(
        doc,
        "Future work may include building a Streamlit interface where users enter nutritional values and receive an "
        "instant predicted nutrition grade. The saved XGBoost model and scaler can be used for this deployment step."
    )

    add_heading(doc, "References", 1)
    refs = [
        "Open Food Facts. Open Food Facts database. Available: https://world.openfoodfacts.org/data",
        "T. Chen and C. Guestrin, XGBoost: A Scalable Tree Boosting System, KDD, 2016.",
        "N. V. Chawla et al., SMOTE: Synthetic Minority Over-sampling Technique, Journal of Artificial Intelligence Research, 2002.",
        "F. Pedregosa et al., Scikit-learn: Machine Learning in Python, Journal of Machine Learning Research, 2011.",
        "M. Abadi et al., TensorFlow: Large-scale Machine Learning on Heterogeneous Systems, 2015.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.add_run(ref)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Nutrition Grade Prediction Detailed Report")

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_report()
    print(OUT_FILE)
